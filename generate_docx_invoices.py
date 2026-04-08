"""
Serverless-compatible DOCX invoice generator.
Generates all employee invoices in a single Word document, all in memory.
"""
import io

# Signature map (same as original)
SIG_MAP = {
    "Ahmed Azam":                   "Ahmed_Azam-removebg-preview.png",
    "Anand Ganesh  Chopade":        "Anand Ganesh Chopade.png",
    "Arundathi Jalagam":            "Arundathi_Jalagam-removebg-preview.png",
    "Bonagiri Rehana":              "Bonagiri_Rehana-removebg-preview.png",
    "Chouta Keerthana":             "Chouta_Keerthana-removebg-preview.png",
    "Deshapaga Raghavendar":        "Deshapaga_Raghavendar-removebg-preview.png",
    "Dhenumakonda Lavanya":         "Dhenumakonda_Lavanya-removebg-preview.png",
    "Diravath Mounika":             "Diravath_Mounika-removebg-preview.png",
    "Gaja Bala Narayana":           "Gaja Bala Narayana-removebg-preview.png",
    "Gopala Saritha":               "Gopala_Saritha-removebg-preview.png",
    "Gudise Hemanth Kumar":         "Gudise_Hemanth_Kumar-removebg-preview.png",
    "Jetti Hima Sindhu":            "Jetti_Hima_Sindhu-removebg-preview.png",
    "K THIRUPATHAIAH":              "K_THIRUPATHAIAH-removebg-preview.png",
    "Kanche Srisailam":             "Kanche_Srisailam-removebg-preview.png",
    "Kasireddy Harish":             "Kasireddy_Harish-removebg-preview.png",
    "Kasturi Sathish":              "Kasturi_Sathish-removebg-preview.png",
    "KATRAVATH MANGESH":            "KATRAVATH_MANGESH-removebg-preview.png",
    "Katravath Mohan Rathod":       "Katravath_Mohan_Rathod-removebg-preview.png",
    "Katravath Radhika":            "Katravath_Radhika-removebg-preview.png",
    "Kodavath Swamy":               "Kodavath_Swamy-removebg-preview.png",
    "KOLLA SUDHAKAR":               "KOLLA_SUDHAKAR-removebg-preview.png",
    "Kunduru Upender Reddy":        "Kunduru_Upender_Reddy-removebg-preview.png",
    "M JATHIN SAI":                 "M_JATHIN_SAI-removebg-preview.png",
    "M Sunitha":                    "M_SUNITHA-removebg-preview.png",
    "M Swarupa":                    "M_Swarupa-removebg-preview.png",
    "Meka Esthar Rani":             "Meka Esthar Rani-removebg-preview.png",
    "Mamidi Raj kumar":             "Mamidi_Raj_kumar-removebg-preview.png",
    "Md Mainoddin":                 "Md_Mainoddin-removebg-preview.png",
    "Mekala Anusha":                "MEKALA_ANUSHA-removebg-preview.png",
    "MOHAMMAD NAZIYA":              "MOHAMMAD_NAZIYA-removebg-preview.png",
    "Mohammad Shaheen Begum":       "Mohammad Shaheen Begum-removebg-preview.png",
    "Mohammed Yaqoob khan":         "Mohammed_Yaqoob_Khan-removebg-preview.png",
    "Mudavath Balakoti":            "Mudavath_Balakoti-removebg-preview.png",
    "MUDAVATH KIRAN":               "MUDAVATH_KIRAN-removebg-preview.png",
    "Mudavath Ramesh":              "Mudavath_Ramesh-removebg-preview.png",
    "Neeli Sreevani":               "Neeli SreevaniSignature remove.png",
    "Padira Radhika":               "Padira_Radhika-removebg-preview.png",
    "Pandiri Punith kumar":         "Pandiri_Punith_kumar-removebg-preview.png",
    "Poojari Ramu":                 "Pujari_Ramu-removebg-preview.png",
    "Porandla Chandar":             "Porandla_Chandar-removebg-preview.png",
    "Ramavath Saimahesh Nayak":     "Ramavath_Saimahesh_Nayak-removebg-preview.png",
    "Ramavath Uma Mahesh":          "Ramavath_Uma_Mahesh-removebg-preview.png",
    "Ranabotu Saidi Reddy":         "Ranabotu_Saidi_Reddy-removebg-preview.png",
    "Sadde Sindhura":               "Sadde_Sindhura-removebg-preview.png",
    "Shaikh abdul Avesh":           "Shaikh_abdul_Avesh-removebg-preview.png",
    "Shivarathri Swapna":           "Shivarathri_Swapna-removebg-preview.png",
    "Tandra Sabastin":              "Tandra_Sabastin-removebg-preview.png",
    "Tabassum Afreen":              "Thabasum_Afreen-removebg-preview.png",
    "Thatipally Manoj Kumar":       "Thatipally_Manoj_Kumar-removebg-preview.png",
    "Thurpati vijay Baskar":        "Thurpati_vijay_Baskar-removebg-preview.png",
    "Ushanolla Ravali":             "Ushanolla_Ravali-removebg-preview.png",
    "Ushanula Ramya":               "Ushanula_Ramya-removebg-preview.png",
    "Chejerla Nagavamsidhar Reddy": "Chejerla_Nagavamsidhar_Reddy-removebg-preview.png",
}

MONTH = "Feb- 2026"
DATE = "28-02-2026"
MONTH_LABEL = "Feb'26"

def rupees(n):
    return f"{int(round(n)):,}".replace(",", ",")

def get_sig_filename(emp):
    """Get signature filename for an employee."""
    name_key = emp.get('service_provider') or emp.get('name', '')
    filename = SIG_MAP.get(name_key)
    if not filename:
        lower_key = name_key.strip().lower()
        for k, v in SIG_MAP.items():
            if k.strip().lower() == lower_key:
                return v
    return emp.get('sig_filename') or filename


def generate_docx_in_memory(employees):
    """
    Generates a DOCX document with all employee invoices in memory.
    Returns a BytesIO buffer containing the DOCX file.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, Twips
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        # Fallback if python-docx is not available
        return _generate_simple_docx(employees)
    
    doc = Document()
    
    # Set page margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    def set_cell_shading(cell, color):
        """Set background color for a table cell."""
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), color)
        cell._tc.get_or_add_tcPr().append(shading)
    
    def add_paragraph(doc, text='', bold=False, size=10, space_after=6, align='left'):
        p = doc.add_paragraph()
        if text:
            run = p.add_run(text)
            run.bold = bold
            run.font.size = Pt(size)
        p.paragraph_format.space_after = Pt(space_after)
        if align == 'center':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == 'right':
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        return p
    
    def add_bold_text_paragraph(doc, parts, space_after=6, align='left'):
        """Add paragraph with mixed bold/normal text."""
        p = doc.add_paragraph()
        for text, is_bold in parts:
            run = p.add_run(text)
            run.bold = is_bold
            run.font.size = Pt(10)
        p.paragraph_format.space_after = Pt(space_after)
        if align == 'center':
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == 'right':
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        return p
    
    for idx, emp in enumerate(employees):
        projs = emp.get('projects', [])
        total_amt = emp.get('total_amount', 0)
        
        # Header info
        add_paragraph(doc, f"Date: - {DATE}", bold=True, space_after=4)
        add_paragraph(doc, "")
        add_paragraph(doc, "TO,", bold=True, space_after=2)
        add_paragraph(doc, "CUBE Highways Technologies Private Limited,", bold=True, space_after=2)
        add_paragraph(doc, "3rd Floor, GMR Aero Towers – 2,", space_after=2)
        add_paragraph(doc, "Mamidipally Village, Saroor Nagar Mandal,", space_after=2)
        add_paragraph(doc, "Ranga Reddy, Hyderabad, Telangana - 500108", space_after=4)
        add_bold_text_paragraph(doc, [("GST No- ", True), ("36AAKCC7533R1ZW", False)], space_after=2)
        add_bold_text_paragraph(doc, [("PAN No- ", True), ("AAKCC7533R", False)], space_after=4)
        add_paragraph(doc, "Sir,", bold=True, space_after=4)
        
        # Subject line
        subject_parts = [
            ("Subject: ", True),
            ("Consultant fee for ", False),
            (MONTH, True),
            (" data processing & Analysis ", False),
            (f"Rs.{rupees(total_amt)}/-", True),
            (" per month. The commercials are mentioned below.", False)
        ]
        add_bold_text_paragraph(doc, subject_parts, space_after=8)
        
        # Fee table
        rate = 16500 / 28
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header row
        hdr_cells = table.rows[0].cells
        headers = ['Particulars', f'No. of working days in {MONTH_LABEL}', 'WBS Elements', 'Payable Amount (Rs.)']
        widths = [Inches(2.5), Inches(1.3), Inches(3), Inches(1.2)]
        
        for i, (cell, header) in enumerate(zip(hdr_cells, headers)):
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_shading(cell, 'BDD7EE')
            cell.width = widths[i]
        
        # Data rows
        if not projs:
            row = table.add_row().cells
            row[0].text = "Consultant fee – Data Processing"
            row[1].text = str(emp.get('attendance', 0))
            row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row[2].text = ""
            row[3].text = rupees(total_amt)
            row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for _ in range(2):
                table.add_row()
        else:
            for i, proj in enumerate(projs):
                row = table.add_row().cells
                row[0].text = "Consultant fee – Data Processing" if i == 0 else ""
                row[1].text = str(proj.get('days', 0))
                row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                row[2].text = proj.get('wbs', '')
                row[3].text = rupees(round(proj.get('days', 0) * rate))
                row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            while len(projs) < 3:
                table.add_row()
                projs.append({})
        
        # Total row
        total_row = table.add_row().cells
        total_row[0].text = ""
        total_row[1].text = str(emp.get('attendance', 0))
        total_row[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        total_row[2].text = "Total Pay"
        total_row[2].paragraphs[0].runs[0].bold = True
        total_row[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        total_row[3].text = f"Rs. {rupees(total_amt)}/-"
        total_row[3].paragraphs[0].runs[0].bold = True
        total_row[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_cell_shading(total_row[2], 'E2EFDA')
        set_cell_shading(total_row[3], 'E2EFDA')
        
        add_paragraph(doc, "", space_after=8)
        add_paragraph(doc, "Thanking you and always assuring you of our best services.", space_after=6)
        add_paragraph(doc, "Yours faithfully", bold=True, space_after=6)
        add_paragraph(doc, "Authorised Signature", bold=True, space_after=20)
        
        # Employee details
        add_bold_text_paragraph(doc, [("Service Provider: ", True), (emp.get('service_provider') or emp.get('name', ''), False)], space_after=2)
        add_bold_text_paragraph(doc, [("Address: ", True), (emp.get('address', ''), False)], space_after=2)
        add_bold_text_paragraph(doc, [("Email- ", True), (emp.get('email', ''), False)], space_after=2)
        add_bold_text_paragraph(doc, [("Contact No. ", True), (emp.get('contact', ''), False)], space_after=2)
        add_bold_text_paragraph(doc, [("PAN No- ", True), (emp.get('pan', ''), False)], space_after=6)
        add_paragraph(doc, "Bank details below:", bold=True, space_after=4)
        
        # Bank table
        bank_table = doc.add_table(rows=2, cols=4)
        bank_table.style = 'Table Grid'
        
        bank_headers = ['Account-Name', 'Bank Name', 'Bank Account Number', 'IFSC Code']
        bank_widths = [Inches(1.8), Inches(1.8), Inches(2), Inches(1.8)]
        
        for i, (cell, header) in enumerate(zip(bank_table.rows[0].cells, bank_headers)):
            cell.text = header
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_shading(cell, 'BDD7EE')
        
        bank_data = bank_table.rows[1].cells
        bank_data[0].text = emp.get('account_name', '')
        bank_data[1].text = emp.get('bank_name', '')
        bank_data[2].text = str(emp.get('account_number', ''))
        bank_data[3].text = emp.get('ifsc', '')
        
        # Add page break except for last employee
        if idx < len(employees) - 1:
            doc.add_page_break()
    
    # Save to BytesIO buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _generate_simple_docx(employees):
    """
    Fallback: Generate a simple text-based document if python-docx is not available.
    This creates a basic document without complex formatting.
    """
    buffer = io.BytesIO()
    
    # Create a minimal ZIP structure for DOCX (Office Open XML)
    import zipfile
    from xml.etree import ElementTree as ET
    
    # This is a simplified fallback - in practice, you should install python-docx
    # For Vercel, add python-docx to your requirements.txt
    
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
        # Minimal DOCX structure
        content_types = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
    <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
    <Default Extension="xml" ContentType="application/xml"/>
    <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>'''
        zf.writestr('[Content_Types].xml', content_types)
        zf.writestr('_rels/.rels', '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
    <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>''')
        
        # Build document content
        doc_content = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
<w:body>'''
        
        for emp in employees:
            total = emp.get('total_amount', 0)
            doc_content += f'''<w:p><w:r><w:t>Date: - {DATE}</w:t></w:r></w:p>
<w:p><w:r><w:t>TO,</w:t></w:r></w:p>
<w:p><w:r><w:t>CUBE Highways Technologies Private Limited</w:t></w:r></w:p>
<w:p><w:r><w:t>Subject: Consultant fee for {MONTH} - Rs.{rupees(total)}/-</w:t></w:r></w:p>
<w:p><w:r><w:t>Service Provider: {emp.get('service_provider') or emp.get('name', 'N/A')}</w:t></w:r></w:p>
<w:p><w:r><w:t>Bank: {emp.get('bank_name', 'N/A')} - A/C: {emp.get('account_number', 'N/A')}</w:t></w:r></w:p>
<w:p><w:r><w:t>IFSC: {emp.get('ifsc', 'N/A')}</w:t></w:r></w:p>
<w:p><w:r><w:t>---</w:t></w:r></w:p>'''
        
        doc_content += '</w:body></w:document>'
        zf.writestr('word/document.xml', doc_content)
    
    buffer.seek(0)
    return buffer
